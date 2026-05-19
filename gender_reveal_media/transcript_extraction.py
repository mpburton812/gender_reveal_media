from __future__ import annotations

import hashlib
import logging
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from urllib.parse import unquote, urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document

logger = logging.getLogger(__name__)

_ALLOWED_SUFFIXES = frozenset({".docx", ".odt", ".doc", ".pdf", ".html"})

_CONTENT_TYPE_TO_EXT: dict[str, str] = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/msword": ".doc",
    "application/vnd.oasis.opendocument.text": ".odt",
    "application/pdf": ".pdf",
    "text/html": ".html",
    "application/xhtml+xml": ".html",
}


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="replace")).hexdigest()


def _filename_from_content_disposition(value: str | None) -> str | None:
    if not value:
        return None
    # filename*=UTF-8''name.ext or filename="name.ext"
    star = re.search(r"filename\*\s*=\s*[^']*'[^']*'([^;]+)", value, re.I)
    if star:
        return unquote(star.group(1).strip().strip('"'))
    plain = re.search(r'filename\s*=\s*"([^"]+)"', value, re.I)
    if plain:
        return plain.group(1).strip()
    plain = re.search(r"filename\s*=\s*([^;\s]+)", value, re.I)
    if plain:
        return plain.group(1).strip().strip('"')
    return None


def infer_transcript_extension(
    url: str,
    content_type: str | None,
    content_disposition: str | None,
) -> str | None:
    """Pick a file suffix from response metadata (after redirects)."""
    fn = _filename_from_content_disposition(content_disposition)
    if fn:
        suf = Path(fn).suffix.lower()
        if suf in _ALLOWED_SUFFIXES:
            return suf

    ct = (content_type or "").split(";", 1)[0].strip().lower()
    mapped = _CONTENT_TYPE_TO_EXT.get(ct)
    if mapped:
        return mapped

    path = urlparse(url).path.split("?", 1)[0].lower()
    for ext in _ALLOWED_SUFFIXES:
        if path.endswith(ext):
            return ext
    return None


def _pandoc_input_format(path: Path) -> str | None:
    """Pandoc input format flag; avoid '-f auto' (not supported on older pandoc, e.g. Ubuntu apt)."""
    return {".docx": "docx", ".odt": "odt", ".doc": "doc"}.get(path.suffix.lower())


def _pandoc_to_plain(path: Path) -> str | None:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return None
    fmt = _pandoc_input_format(path)
    if not fmt:
        return None
    try:
        proc = subprocess.run(
            [pandoc, "-f", fmt, "-t", "plain", str(path)],
            check=False,
            capture_output=True,
            text=True,
            timeout=300,
        )
        if proc.returncode != 0:
            logger.warning("pandoc failed rc=%s stderr=%s", proc.returncode, proc.stderr[:500])
            return None
        return proc.stdout.strip()
    except subprocess.TimeoutExpired:
        logger.warning("pandoc timed out for %s", path)
        return None


def _extract_doc_antiword(path: Path) -> str | None:
    antiword = shutil.which("antiword")
    if not antiword:
        return None
    try:
        proc = subprocess.run(
            [antiword, str(path)],
            check=False,
            capture_output=True,
            text=True,
            timeout=120,
        )
        if proc.returncode != 0:
            logger.warning("antiword failed rc=%s stderr=%s", proc.returncode, proc.stderr[:500])
            return None
        return proc.stdout.strip()
    except subprocess.TimeoutExpired:
        logger.warning("antiword timed out for %s", path)
        return None


def _convert_doc_to_docx(path: Path, workdir: Path) -> Path | None:
    office = shutil.which("soffice") or shutil.which("libreoffice")
    if not office:
        return None
    try:
        proc = subprocess.run(
            [
                office,
                "--headless",
                "--norestore",
                "--convert-to",
                "docx",
                "--outdir",
                str(workdir),
                str(path),
            ],
            check=False,
            capture_output=True,
            text=True,
            timeout=300,
        )
        if proc.returncode != 0:
            logger.warning("libreoffice convert failed rc=%s stderr=%s", proc.returncode, proc.stderr[:500])
            return None
        out = workdir / f"{path.stem}.docx"
        return out if out.is_file() else None
    except subprocess.TimeoutExpired:
        logger.warning("libreoffice timed out for %s", path)
        return None


def _extract_docx_python(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)
    return "\n".join(parts).strip()


def _extract_odt_python(path: Path) -> str:
    from odf import teletype
    from odf.opendocument import load
    from odf import text as odf_text

    doc = load(str(path))
    chunks: list[str] = []
    for el in doc.getElementsByType(odf_text.P):
        t = teletype.extractText(el).strip()
        if t:
            chunks.append(t)
    return "\n".join(chunks).strip()


def _extract_pdf_python(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        t = (page.extract_text() or "").strip()
        if t:
            parts.append(t)
    return "\n\n".join(parts).strip()


def _extract_html_transcript(html: str) -> str:
    soup = BeautifulSoup(html, "lxml")
    root = soup.find("article") or soup.find("main") or soup.body
    if not root:
        raise RuntimeError("HTML transcript page has no article/main body")
    skip_prefixes = (
        "share this",
        "support team",
        "share on ",
        "listen on ",
    )
    parts: list[str] = []
    for el in root.find_all(["p", "h1", "h2", "h3", "li"]):
        t = el.get_text(" ", strip=True)
        if not t or len(t) < 20:
            continue
        low = t.lower()
        if any(low.startswith(p) for p in skip_prefixes):
            continue
        parts.append(t)
    text = "\n\n".join(parts).strip()
    if not text:
        text = root.get_text("\n", strip=True)
    return text


def extract_local_file(path: Path, *, workdir: Path | None = None) -> str:
    suffix = path.suffix.lower()
    plain = _pandoc_to_plain(path)
    if plain:
        return plain
    if suffix == ".docx":
        return _extract_docx_python(path)
    if suffix == ".odt":
        return _extract_odt_python(path)
    if suffix == ".pdf":
        return _extract_pdf_python(path)
    if suffix == ".doc":
        plain = _extract_doc_antiword(path)
        if plain:
            return plain
        wd = workdir or path.parent
        converted = _convert_doc_to_docx(path, wd)
        if converted:
            plain = _pandoc_to_plain(converted) or _extract_docx_python(converted)
            if plain:
                return plain
        raise RuntimeError(
            "Legacy .doc could not be converted. Install antiword or LibreOffice in CI."
        )
    raise RuntimeError(f"Unsupported transcript file type: {suffix}")


def download_transcript_text(
    url: str,
    *,
    user_agent: str,
    timeout: int = 120,
) -> tuple[str, str]:
    """
    Download transcript file and return (plain_text, sha256_hex of raw bytes).
    Follows redirects; detects type from Content-Type / Content-Disposition / URL.
    """
    headers = {"User-Agent": user_agent}
    r = requests.get(url, headers=headers, timeout=timeout, allow_redirects=True)
    r.raise_for_status()
    data = r.content
    ext = infer_transcript_extension(
        r.url,
        r.headers.get("content-type"),
        r.headers.get("content-disposition"),
    )
    if not ext:
        ct = r.headers.get("content-type", "")
        raise RuntimeError(
            f"UNSUPPORTED_TRANSCRIPT_URL: cannot detect file type for {url!r} "
            f"(final={r.url!r}, content-type={ct!r})"
        )

    if ext == ".html":
        text = _extract_html_transcript(r.text)
    else:
        tmpdir = Path(tempfile.mkdtemp(prefix="grm-transcript-"))
        try:
            local = tmpdir / f"transcript{ext}"
            local.write_bytes(data)
            text = extract_local_file(local, workdir=tmpdir)
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)

    if not text.strip():
        raise RuntimeError("EMPTY_TRANSCRIPT_AFTER_EXTRACTION")
    digest = hashlib.sha256(data).hexdigest()
    return text, digest
